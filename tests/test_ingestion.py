"""
Unit tests for Step 1: Input Ingestion module.
"""

from __future__ import annotations

import os
import sys

import pytest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from agent.ingestion import (
    extract_docx_text,
    extract_pdf_text,
    extract_raw_text,
    ingest_inputs,
    validate_file,
    MAX_FILE_SIZE_BYTES,
)


SAMPLE_DIR = os.path.join(os.path.dirname(__file__), "sample_data")


class TestValidateFile:
    """Tests for file validation."""
    
    def test_valid_pdf_file(self):
        """Test that a valid PDF file passes validation."""
        valid, error = validate_file(b"%PDF-1.4 test content", "resume.pdf")
        assert valid is True
        assert error == ""
    
    def test_valid_docx_file(self):
        """Test that a valid DOCX file passes validation."""
        valid, error = validate_file(b"PK\x03\x04 test content", "resume.docx")
        assert valid is True
        assert error == ""
    
    def test_file_size_exceeded(self):
        """Test that files over 10 MB are rejected."""
        large_content = b"x" * (MAX_FILE_SIZE_BYTES + 1)
        valid, error = validate_file(large_content, "large.pdf")
        assert valid is False
        assert "exceeds maximum size" in error
    
    def test_unsupported_extension(self):
        """Test that unsupported file extensions are rejected."""
        valid, error = validate_file(b"some content", "malware.exe")
        assert valid is False
        assert "unsupported extension" in error
        
        valid2, error2 = validate_file(b"some content", "archive.zip")
        assert valid2 is False
        assert "unsupported extension" in error2
    
    def test_valid_txt_file(self):
        """Test that .txt files pass validation."""
        valid, error = validate_file(b"This is a JD text file.", "jd.txt")
        assert valid is True
        assert error == ""
    
    def test_valid_json_file(self):
        """Test that .json files pass validation."""
        valid, error = validate_file(b'{"key": "value"}', "profile.json")
        assert valid is True
        assert error == ""


class TestExtractPdfText:
    """Tests for PDF text extraction."""
    
    def test_extract_sample_pdf(self):
        """Test extracting text from a sample PDF file."""
        pdf_path = os.path.join(SAMPLE_DIR, "resume_strong_match.pdf")
        if os.path.exists(pdf_path):
            with open(pdf_path, "rb") as f:
                content = f.read()
            text = extract_pdf_text(content)
            assert len(text) > 100
            assert "ALEXANDRA" in text or "MARTINEZ" in text or "Python" in text
    
    def test_extract_invalid_pdf(self):
        """Test handling of invalid PDF content."""
        text = extract_pdf_text(b"not a pdf file")
        assert text == "" or isinstance(text, str)


class TestExtractDocxText:
    """Tests for DOCX text extraction."""
    
    def test_extract_docx(self):
        """Test extracting text from a DOCX file."""
        # Create a minimal docx for testing
        try:
            from docx import Document
            doc = Document()
            doc.add_paragraph("This is a test document.")
            doc.add_paragraph("Python Django React")
            
            import io
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            
            text = extract_docx_text(buf.read())
            assert "test document" in text
            assert "Python" in text
        except ImportError:
            pytest.skip("python-docx not available")


class TestIngestInputs:
    """Tests for the main ingestion function."""
    
    def test_jd_text_only_no_candidates(self):
        """Test that missing candidates raises an error."""
        result = ingest_inputs(jd_text="Senior Engineer position...")
        assert len(result.errors) > 0
        assert "No candidate inputs" in result.errors[0]
    
    def test_no_jd(self):
        """Test that missing JD raises an error."""
        result = ingest_inputs(resume_files=[(b"%PDF test", "resume.pdf")])
        assert len(result.errors) > 0
        assert "No job description" in result.errors[0]
    
    def test_jd_text_with_resume(self):
        """Test successful ingestion with JD text and resume."""
        result = ingest_inputs(
            jd_text="Senior Python Engineer position. Must know Django and React.",
            resume_files=[(b"%PDF test resume content", "candidate.pdf")],
        )
        assert len(result.errors) == 0
        assert result.jd_raw == "Senior Python Engineer position. Must know Django and React."
        assert len(result.candidates) == 1
        assert result.candidates[0]["id"] == "cand_001"
        assert result.candidates[0]["source"] == "resume_pdf"
    
    def test_multiple_inputs(self):
        """Test ingestion with multiple resume and LinkedIn inputs."""
        result = ingest_inputs(
            jd_text="Engineer position",
            resume_files=[
                (b"%PDF resume 1", "candidate1.pdf"),
                (b"%PDF resume 2", "candidate2.pdf"),
            ],
            linkedin_urls=["https://linkedin.com/in/person1"],
        )
        assert len(result.errors) == 0
        assert len(result.candidates) == 3
        assert result.candidates[0]["source"] == "resume_pdf"
        assert result.candidates[2]["source"] == "linkedin_url"
