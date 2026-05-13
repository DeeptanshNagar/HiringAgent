"""
Step 1: Input Ingestion Module.
Handles file upload validation, text extraction, and structured output.
Supports PDF, DOCX, TXT, JSON (LinkedIn export) files.
"""

from __future__ import annotations

import io
import json
import logging
import mimetypes
import os
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)

# File size limit: 10 MB
MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024

# Supported MIME types
SUPPORTED_MIME_TYPES = {
    "application/pdf": [".pdf"],
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": [".docx"],
    "application/json": [".json"],
    "text/plain": [".txt"],
}

# All supported extensions
SUPPORTED_EXTENSIONS = set()
for exts in SUPPORTED_MIME_TYPES.values():
    SUPPORTED_EXTENSIONS.update(exts)


def validate_file(file_bytes: bytes, filename: str) -> Tuple[bool, str]:
    """
    Validate a file's size and extension.
    
    Args:
        file_bytes: Raw file content.
        filename: Original filename with extension.
        
    Returns:
        Tuple of (is_valid, error_message).
    """
    # Check file size
    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        return False, (
            f"File '{filename}' exceeds maximum size of 10 MB "
            f"({len(file_bytes) / (1024*1024):.1f} MB)"
        )
    
    # Check extension
    _, ext = os.path.splitext(filename.lower())
    if ext not in SUPPORTED_EXTENSIONS:
        return False, (
            f"File '{filename}' has unsupported extension '{ext}'. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )
    
    return True, ""


def extract_pdf_text(file_bytes: bytes) -> str:
    """
    Extract text from PDF using PyMuPDF (fitz) with pdfplumber fallback.
    
    Args:
        file_bytes: PDF file content.
        
    Returns:
        Extracted text string.
    """
    text_parts: list[str] = []
    
    # Try PyMuPDF first
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(page_text)
        doc.close()
        
        combined = "\n".join(text_parts)
        if len(combined.strip()) > 50:
            return combined
    except Exception as e:
        logger.warning(f"PyMuPDF extraction failed: {e}, trying pdfplumber fallback")
    
    # Fallback to pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text)
        
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"pdfplumber extraction also failed: {e}")
        return ""


def extract_docx_text(file_bytes: bytes) -> str:
    """
    Extract text from DOCX file using python-docx.
    
    Args:
        file_bytes: DOCX file content.
        
    Returns:
        Extracted text string.
    """
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        
        text_parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ""


def extract_text_file(file_bytes: bytes) -> str:
    """
    Decode plain text file.
    
    Args:
        file_bytes: Text file content.
        
    Returns:
        Decoded text string.
    """
    for encoding in ["utf-8", "latin-1", "cp1252"]:
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return ""


def parse_linkedin_json(file_bytes: bytes) -> Dict[str, Any]:
    """
    Parse a LinkedIn JSON export file.
    
    Args:
        file_bytes: JSON file content.
        
    Returns:
        Parsed JSON dict.
    """
    try:
        text = file_bytes.decode("utf-8")
        return json.loads(text)
    except (UnicodeDecodeError, json.JSONDecodeError) as e:
        logger.error(f"LinkedIn JSON parsing failed: {e}")
        return {}


def extract_raw_text(file_bytes: bytes, filename: str) -> str:
    """
    Extract raw text from a file based on its extension.
    
    Args:
        file_bytes: File content.
        filename: Original filename.
        
    Returns:
        Extracted text.
    """
    _, ext = os.path.splitext(filename.lower())
    
    if ext == ".pdf":
        return extract_pdf_text(file_bytes)
    elif ext == ".docx":
        return extract_docx_text(file_bytes)
    elif ext == ".txt":
        return extract_text_file(file_bytes)
    elif ext == ".json":
        data = parse_linkedin_json(file_bytes)
        # For LinkedIn JSON, we need special handling - return as string for now
        return json.dumps(data, indent=2)
    else:
        logger.warning(f"Unknown extension '{ext}' for file '{filename}', trying text extraction")
        return extract_text_file(file_bytes)


@dataclass
class IngestionResult:
    """Result of the ingestion step."""
    jd_raw: str = ""
    candidates: List[Dict[str, Any]] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)
    files_processed: List[str] = field(default_factory=list)


def ingest_inputs(
    jd_text: Optional[str] = None,
    jd_file: Optional[Tuple[bytes, str]] = None,
    resume_files: Optional[List[Tuple[bytes, str]]] = None,
    linkedin_json_files: Optional[List[Tuple[bytes, str]]] = None,
    linkedin_urls: Optional[List[str]] = None,
) -> IngestionResult:
    """
    Main ingestion function that processes all inputs.
    
    Args:
        jd_text: Plain text job description.
        jd_file: Tuple of (file_bytes, filename) for JD.
        resume_files: List of (file_bytes, filename) for resumes.
        linkedin_json_files: List of (file_bytes, filename) for LinkedIn JSON exports.
        linkedin_urls: List of LinkedIn profile URLs.
        
    Returns:
        IngestionResult with structured data.
    """
    result = IngestionResult()
    
    # Process JD
    if jd_text:
        result.jd_raw = jd_text
        result.files_processed.append("jd_text_input")
    elif jd_file:
        file_bytes, filename = jd_file
        is_valid, error = validate_file(file_bytes, filename)
        if not is_valid:
            result.errors.append(f"JD validation failed: {error}")
            return result
        result.jd_raw = extract_raw_text(file_bytes, filename)
        result.files_processed.append(filename)
    else:
        result.errors.append("No job description provided. Please provide JD text or upload a JD file.")
        return result
    
    if not result.jd_raw or not result.jd_raw.strip():
        result.errors.append("Job description is empty after extraction.")
        return result
    
    # Process resume files
    candidate_id = 0
    if resume_files:
        for file_bytes, filename in resume_files:
            candidate_id += 1
            is_valid, error = validate_file(file_bytes, filename)
            if not is_valid:
                result.errors.append(f"Resume '{filename}': {error}")
                continue
            
            raw_text = extract_raw_text(file_bytes, filename)
            _, ext = os.path.splitext(filename.lower())
            
            source = "resume_pdf" if ext == ".pdf" else "resume_docx" if ext == ".docx" else "resume"
            
            result.candidates.append({
                "id": f"cand_{candidate_id:03d}",
                "source": source,
                "filename": filename,
                "raw_text": raw_text,
            })
            result.files_processed.append(filename)
    
    # Process LinkedIn JSON files
    if linkedin_json_files:
        for file_bytes, filename in linkedin_json_files:
            candidate_id += 1
            is_valid, error = validate_file(file_bytes, filename)
            if not is_valid:
                result.errors.append(f"LinkedIn JSON '{filename}': {error}")
                continue
            
            raw_text = extract_raw_text(file_bytes, filename)
            
            result.candidates.append({
                "id": f"cand_{candidate_id:03d}",
                "source": "linkedin_json",
                "filename": filename,
                "raw_text": raw_text,
            })
            result.files_processed.append(filename)
    
    # Process LinkedIn URLs
    if linkedin_urls:
        for url in linkedin_urls:
            if url and url.strip():
                candidate_id += 1
                result.candidates.append({
                    "id": f"cand_{candidate_id:03d}",
                    "source": "linkedin_url",
                    "url": url.strip(),
                    "raw_text": "",  # Will be fetched later
                })
                result.files_processed.append(f"linkedin_url_{candidate_id}")
    
    # Validation: need at least 1 candidate
    if not result.candidates:
        result.errors.append("No candidate inputs provided. Please upload at least one resume or LinkedIn profile.")
    
    logger.info(
        f"Ingestion complete: {len(result.candidates)} candidates, "
        f"JD length: {len(result.jd_raw)} chars, "
        f"files: {result.files_processed}"
    )
    
    return result
